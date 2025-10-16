import os
import json
import shutil
from dotenv import load_dotenv
from django.conf import settings
from django.core.files.storage import default_storage
from django.http import JsonResponse, HttpResponseBadRequest, StreamingHttpResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt

from .models import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status
from django.shortcuts import get_object_or_404


from .serializers import DocumentSerializer
from cryptography.fernet import InvalidToken

from django.utils.dateparse import parse_date
from cryptography.fernet import Fernet
from PyPDF2 import PdfReader, PdfWriter
import logging
from .logger import log_exception, log_exceptions
import uuid
import time
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from io import BytesIO
from bs4 import BeautifulSoup

# Using single HTML-only LLM call for uploads and full-document processing

# Setup logger
logger = logging.getLogger(__name__)

# Updated import with new streaming function
from .vertex_model import call_gemini_api_with_streaming, MODEL_ID
import yaml
import configparser
import io
from PIL import Image

def make_llm_call(prompt_text, input_data, response_mime_type, max_pages, progress_callback, call_type):
    """Helper function to make a single LLM call"""
    try:
        progress_callback(f"Starting {call_type} generation...")
        start_time = time.time()

        response = call_gemini_api_with_streaming(
            prompt_text=prompt_text,
            input_data=input_data,
            response_mime_type=response_mime_type,
            max_pages=max_pages,
            progress_callback=progress_callback
        )
        response_time = time.time() - start_time
        progress_callback(f"Completed {call_type} generation in {response_time:.2f}s")

        return response, None  # Return (response, error)
    except Exception as e:
        error_msg = f"Error in {call_type} generation: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return None, error_msg

# Initialize prompts
APP_CONFIG = {}
CONFIG_FILE_PATH = os.path.join(settings.BASE_DIR, 'Prompts', 'prompts.yaml')

def load_prompts():
    """Load prompts from YAML file"""
    global APP_CONFIG
    try:
        with open(CONFIG_FILE_PATH, 'r', encoding='utf-8') as f:
            APP_CONFIG = yaml.safe_load(f)
        logger.info(f"Successfully loaded prompts from {CONFIG_FILE_PATH}")
        logger.debug(f"APP_CONFIG keys: {list(APP_CONFIG.keys())}")
        if 'prompts' in APP_CONFIG and isinstance(APP_CONFIG['prompts'], dict):
            logger.debug(f"prompts keys: {list(APP_CONFIG['prompts'].keys())}")
    except FileNotFoundError:
        logger.critical(f"Prompts file not found at {CONFIG_FILE_PATH}. Please ensure it exists.")
        raise 
    except Exception as e:
        logger.critical(f"Error loading prompts from YAML at {CONFIG_FILE_PATH}: {e}", exc_info=True)
        raise

load_prompts()

env_path = os.path.join(settings.BASE_DIR, '.env') if hasattr(settings, 'BASE_DIR') else None
if env_path and os.path.exists(env_path):
    load_dotenv(env_path)
else:
    load_dotenv()

def safe_json_load(raw_string: str):
    if not raw_string or not raw_string.strip():
        raise json.JSONDecodeError("Empty or whitespace-only string", raw_string, 0)

    cleaned = raw_string.strip()

    if cleaned.startswith("```json"):
        lines = cleaned.splitlines()
        if len(lines) > 2 and lines[-1].strip() == "```":
            cleaned = "\n".join(lines[1:-1])
        else:
            cleaned = "\n".join(lines[1:])

    return json.loads(cleaned)

def get_fernet_key():
    try:
        config = configparser.ConfigParser()
        config_path = os.path.join(os.path.dirname(__file__), 'config.properties')
        
        if not os.path.exists(config_path):
            logger.error(f"Config file not found at {config_path}")
            return None
        
        config.read(config_path)
        key = config.get('Input', 'FERNET_KEY')
        if not key:
            logger.error("Fernet key not found in config.properties")
            return None
        try:
            Fernet(key.encode())
        except Exception as e:
            logger.error(f"Invalid Fernet key: {str(e)}")
            return None
        return key
    except Exception as e:
        logger.error(f"Error loading Fernet key: {str(e)}")
        return None

FERNET_KEY = get_fernet_key()

if not FERNET_KEY:
    logger.error("Fernet key initialization failed")

def encrypt_id(id: int) -> str:
    fernet = Fernet(FERNET_KEY)
    id_bytes = str(id).encode()
    encrypted = fernet.encrypt(id_bytes)
    return encrypted.decode()

def decrypt_id(token: str) -> int:
    fernet = Fernet(FERNET_KEY)
    decrypted = fernet.decrypt(token.encode())
    return int(decrypted.decode())

# Updated views with user restrictions and streaming

class GetDocumentByIdView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, doc_id):
        try:
            logger.info(f"GetDocumentByIdView called with encrypted ID: {doc_id}")

            decrypted_id = decrypt_id(doc_id)
            logger.debug(f"Decrypted ID: {decrypted_id}")

            with log_exceptions(logger):
                doc = get_object_or_404(Document, id=decrypted_id)

                # Updated permission checks with user types
                user = request.user
                is_admin = user.user_type == "admin"

                if not is_admin and doc.userid_id != request.user.id:
                    logger.warning(
                        f"User {request.user.id} attempted to access document {decrypted_id} without permission"
                    )
                    return Response(
                        {"error": "You do not have permission to view this document."},
                        status=status.HTTP_403_FORBIDDEN,
                    )

                logger.info(f"Document {decrypted_id} retrieved successfully")

                file_url = request.build_absolute_uri(doc.file.url)

                response_data = {
                    "status": "success",
                    "filePath": file_url,
                    "json_data": doc.json_data,
                    "html_data": doc.html_data,
                    "input_token": doc.input_token,
                    "output_token": doc.output_token,
                    "api_response_time": doc.api_response_time,
                    "db_save_time": doc.db_save_time,
                    "llm_model_used": doc.llm_model_used,
                    "pages_processed": getattr(doc, 'pages_processed', 1),
                    "is_full_document": getattr(doc, 'is_full_document', False),
                }

            return Response(response_data, status=status.HTTP_200_OK)

        except InvalidToken:
            logger.error("Invalid or corrupted document ID provided.", exc_info=True)
            return Response(
                {"error": "Invalid or corrupted document ID"},
                status=status.HTTP_400_BAD_REQUEST
            )
        except Exception:
            logger.error("Error while fetching document by ID", exc_info=True)
            return Response(
                {"error": "An internal server error occurred while retrieving the document."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

class UserDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            user = request.user
            logger.info(f"UserDocumentView accessed by user ID: {user.id}")

            # Determine if the requesting user is an admin based on user_type
            is_admin = user.user_type == "admin"

            if is_admin:
                documents = Document.objects.all()
                logger.info("Admin user detected: Fetching all documents.")
            else:
                documents = Document.objects.filter(userid=user)
                logger.info(f"Fetching documents for user ID: {user.id}")

            serializer = DocumentSerializer(documents, many=True)
            serialized_data = serializer.data
            logger.info(f"{len(serialized_data)} documents retrieved successfully.")

            # Encrypt the 'id' field and ensure html_data is present to be consistent with UserDocumentView
            for idx, doc in enumerate(serialized_data):
                try:
                    doc['id'] = encrypt_id(doc['id'])
                except Exception:
                    logger.exception(f"Failed to encrypt document id for doc: {doc.get('id')}")
                # Ensure html_data exists in the serialized output (copy from model object if needed)
                try:
                    model_doc = documents[idx]
                    doc['html_data'] = getattr(model_doc, 'html_data', None)
                except Exception:
                    # leave html_data as-is if any issue
                    pass

            total_input_tokens = sum(getattr(doc, "input_token", 0) or 0 for doc in documents)
            total_output_tokens = sum(getattr(doc, "output_token", 0) or 0 for doc in documents)

            # Add user usage information
            usage_info = user.get_usage_info() if hasattr(user, 'get_usage_info') else {}
            
            logger.info(f"{len(serialized_data)} documents retrieved. Total input: {total_input_tokens}, output: {total_output_tokens}")
            
            return Response({
                "count": documents.count(),
                "documents": serialized_data,
                "total_input_tokens": total_input_tokens,
                "total_output_tokens": total_output_tokens,
                "user_usage": usage_info,  # Add usage information
            }, status=status.HTTP_200_OK)
            
        except Exception:
            logger.error("Exception occurred while fetching user documents.", exc_info=True)
            log_exception(logger)
            return Response({
                "status": "error",
                "message": "An error occurred while retrieving documents. Please try again later."
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class FilteredDocumentView(APIView):
    def post(self, request):
        try:
            user_id = request.data.get('userid')
            date_str = request.data.get('date')

            logger.info(f"FilteredDocumentView called with user_id={user_id} and date={date_str}")

            if not user_id or not date_str:
                logger.warning("Missing 'userid' or 'date' in request body.")
                return Response({
                    "error": "Both 'userid' and 'date' fields are required in the request body."
                }, status=status.HTTP_400_BAD_REQUEST)

            entry_date = parse_date(date_str)
            if not entry_date:
                logger.warning(f"Invalid date format received: {date_str}")
                return Response({
                    "error": "Invalid date format. Please use ISO-MM-DD."
                }, status=status.HTTP_400_BAD_REQUEST)

            # Updated admin check
            user = request.user
            is_admin = user.user_type == "admin"

            try:
                user_id_int = int(user_id)
            except (TypeError, ValueError):
                logger.warning(f"Invalid user ID format received: {user_id}")
                return Response(
                    {"error": "Invalid user ID. It must be an integer."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            if not is_admin and user_id_int != request.user.id:
                logger.warning(
                    f"User {request.user.id} attempted to filter documents for user {user_id} without permission"
                )
                return Response(
                    {"error": "You do not have permission to view these documents."},
                    status=status.HTTP_403_FORBIDDEN,
                )

            documents = Document.objects.filter(userid=user_id_int, entry_date=entry_date)
            serializer = DocumentSerializer(documents, many=True)
            serialized_data = serializer.data

            # Encrypt ids and ensure html_data present in the serialized output
            for idx, doc in enumerate(serialized_data):
                try:
                    doc['id'] = encrypt_id(doc['id'])
                except Exception:
                    logger.exception(f"Failed to encrypt document id for doc: {doc.get('id')}")
                try:
                    model_doc = documents[idx]
                    doc['html_data'] = getattr(model_doc, 'html_data', None)
                except Exception:
                    pass

            logger.info(f"{documents.count()} documents found for user_id={user_id} on {entry_date}")

            return Response({
                "count": documents.count(),
                "documents": serialized_data
            }, status=status.HTTP_200_OK)

        except Exception:
            logger.error("Exception occurred while filtering documents.", exc_info=True)
            log_exception(logger)
            return Response({
                "error": "An internal error occurred. Please try again later."
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class UploadAndProcessFileView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        uploaded_file = request.FILES.get("pdf_file")
        prompt_text_from_request = request.POST.get("prompt_text")
        doc_type = request.POST.get("doc_type")
        process_full_document = request.POST.get("process_full_document", "false").lower() == "true"

        logger.info("Upload request received")

        user = request.user
        input_tokens = 0
        output_tokens = 0

        if not uploaded_file:
            logger.error("Upload failed: 'pdf_file' is missing in the request.", exc_info=True)
            return Response(
                {"status": "error", "message": "Missing 'pdf_file"},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check user permissions and document limits
        can_process, limit_message = user.can_process_document()
        if not can_process:
            logger.warning(f"User {user.id} hit document limit: {limit_message}")
            return Response({
                "status": "error", 
                "message": limit_message,
                "usage_info": user.get_usage_info()
            }, status=status.HTTP_403_FORBIDDEN)

        # Determine page limit based on user type
        max_pages = 3
        if user.user_type in ['power', 'admin'] and process_full_document:
            max_pages = None
        
        # Determine prompt text
        prompt_text = None
        if prompt_text_from_request:
            prompt_text = prompt_text_from_request
        else:
            if doc_type == 'docextraction':
                prompt_text = APP_CONFIG.get('prompts', {}).get('html_generation_prompt', '')
            elif doc_type == 'Bill Reimbursment':
                prompt_text = APP_CONFIG.get('prompts', {}).get('reimbursement_extraction_prompt', '')
            else:
                prompt_text = APP_CONFIG.get('prompts', {}).get('doc_extraction_prompt', '')

        if not prompt_text:
            logger.error(f"Prompt text is empty or not found in prompts.yaml for doc_type: {doc_type}.")
            return Response(
                {"status": "error", "message": "Prompt text could not be determined for the document type."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

        try:
            with log_exceptions(logger):
                # Save uploaded file
                custom_dir = "uploads/pdf_files"
                save_dir = os.path.join(settings.MEDIA_ROOT, custom_dir)
                os.makedirs(save_dir, exist_ok=True)

                file_name = uploaded_file.name
                name_without_ext, extension = os.path.splitext(file_name)
                extension = extension.lower()
                sanitized_name = name_without_ext.replace("/", "_").replace("\\", "_")
                unique_name = f"{sanitized_name}_{uuid.uuid4()}{extension}"
                relative_path = default_storage.save(
                    os.path.join(custom_dir, unique_name), uploaded_file
                )
                absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)

                if extension not in [".jpg", ".jpeg", ".png", ".pdf"]:
                    logger.error("Unsupported file type provided.", exc_info=True)
                    return Response(
                        {"status": "error", "message": "Unsupported file type"},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                # Progress tracking for streaming updates
                progress_messages = []
                
                def progress_callback(message):
                    progress_messages.append({
                        "timestamp": time.time(),
                        "message": message
                    })

                # Extract structured data using a single HTML-only call (no threading, no JSON)
                try:
                    progress_callback("Starting HTML generation...")
                    overall_start = time.time()

                    # Use HTML prompt for upload processing
                    html_prompt_text = APP_CONFIG.get('prompts', {}).get('html_generation_prompt', '')
                    if not html_prompt_text:
                        logger.error("HTML prompt not configured for uploads.")
                        return Response({"error": "HTML prompt not configured"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                    html_response, error = make_llm_call(
                        html_prompt_text,
                        absolute_path,
                        "text/plain",
                        max_pages,
                        progress_callback,
                        "HTML"
                    )

                    overall_time = time.time() - overall_start
                    progress_callback(f"HTML generation completed in {overall_time:.2f}s")

                    if error or not html_response:
                        logger.error(f"HTML generation failed: {error}")
                        return Response({"error": f"Failed to generate HTML: {error}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                    if 'candidates' not in html_response:
                        logger.error("Invalid HTML response format from API.", exc_info=True)
                        return Response({"error": "Invalid HTML response format"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                    # Extract HTML content and pages processed (if provided)
                    try:
                        html_content = html_response['candidates'][0]['content']['parts'][0]['text']
                        html_data = html_content.strip() if html_content and html_content.strip() else None
                        pages_processed = html_response.get('pagesProcessed', 1) if isinstance(html_response, dict) else 1
                        logger.debug("Successfully generated HTML response")
                    except Exception as e:
                        logger.error(f"Error extracting HTML from response: {str(e)}", exc_info=True)
                        return Response({"error": "Error processing HTML response"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                except Exception as e:
                    logger.error(f"Error during HTML generation call: {str(e)}", exc_info=True)
                    return Response({"error": f"Error during HTML generation: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                # Calculate total tokens based on HTML generation usage metadata
                total_input_tokens = 0
                total_output_tokens = 0

                if html_data and html_response and isinstance(html_response, dict) and 'usageMetadata' in html_response:
                    html_usage = html_response['usageMetadata']
                    total_input_tokens += html_usage.get('promptTokenCount', 0)
                    total_output_tokens += html_usage.get('candidatesTokenCount', 0)
                    logger.info(
                        f"HTML Generation - Input Tokens: {html_usage.get('promptTokenCount', 0)}, Output Tokens: {html_usage.get('candidatesTokenCount', 0)}"
                    )

                input_tokens = total_input_tokens
                output_tokens = total_output_tokens

                # Ensure api_response_time is defined. Prefer previously computed overall_time
                try:
                    api_response_time = overall_time
                except NameError:
                    try:
                        api_response_time = time.time() - overall_start
                    except NameError:
                        api_response_time = None

                db_start = time.time()
                doc = Document.objects.create(
                    file_path=relative_path,
                    file=relative_path,
                    html_data=html_data,
                    userid=user,
                    document_type=doc_type,
                    input_token=input_tokens,
                    output_token=output_tokens,
                    api_response_time=api_response_time,
                    llm_model_used=MODEL_ID,
                    pages_processed=pages_processed,
                    is_full_document=process_full_document and user.user_type in ['power', 'admin']
                )
                db_save_time = time.time() - db_start
                doc.db_save_time = db_save_time
                doc.save(update_fields=["db_save_time"])

                # Update user usage counters
                user.increment_usage(pages_processed)

                encrypted_doc_id = encrypt_id(doc.id)
                logger.info(
                    f"Document processed and saved successfully. Document ID: {encrypted_doc_id}"
                )

                # Prepare response with usage information
                response_data = {
                    "status": "success",
                    "document_id": encrypted_doc_id,
                    "pages_processed": pages_processed,
                    "is_full_document": doc.is_full_document,
                    "html_generated": html_data is not None,
                    "progress_messages": progress_messages,
                    "usage_info": user.get_usage_info()
                }

                # Add "Load Full Document" option for power and admin users
                if (user.user_type in ['power', 'admin'] and not process_full_document and
                    pages_processed == 3):  # Only if we actually limited to 3 pages
                    response_data["can_load_full_document"] = True
                    response_data["message"] = "Processed first 3 pages. You can load the full document if needed."

                return Response(response_data, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Error in UploadAndProcessFileView: {str(e)}", exc_info=True)
            return Response(
                {"status": "error", "message": f"An internal server error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

# New view for power users to process full documents
class ProcessFullDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        document_id = request.data.get("document_id")
        
        if not document_id:
            return Response(
                {"status": "error", "message": "Missing document_id"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        user = request.user
        
        # Check if user is power user or admin
        if user.user_type not in ['power', 'admin']:
            return Response(
                {"status": "error", "message": "Only power users and admins can process full documents"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            decrypted_id = decrypt_id(document_id)
            doc = get_object_or_404(Document, id=decrypted_id)
            
            # Check ownership
            is_admin = user.user_type == 'admin'
            if not is_admin and doc.userid_id != user.id:
                return Response(
                    {"error": "You do not have permission to process this document."},
                    status=status.HTTP_403_FORBIDDEN,
                )
            
            # Check if already processed as full document
            if doc.is_full_document:
                return Response(
                    {"status": "error", "message": "Document already processed as full document"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Get the original file path
            absolute_path = os.path.join(settings.MEDIA_ROOT, doc.file_path)
            
            if not os.path.exists(absolute_path):
                return Response(
                    {"status": "error", "message": "Original file not found"},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Determine prompt text based on document type
            prompt_text = None
            if doc.document_type == 'docextraction':
                prompt_text = APP_CONFIG.get('prompts', {}).get('html_generation_prompt', '')
            elif doc.document_type == 'Bill Reimbursment':
                prompt_text = APP_CONFIG.get('prompts', {}).get('reimbursement_extraction_prompt', '')
            else:
                prompt_text = APP_CONFIG.get('prompts', {}).get('doc_extraction_prompt', '')
            
            if not prompt_text:
                return Response(
                    {"status": "error", "message": "Could not determine prompt for document type"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # Progress tracking
            progress_messages = []
            
            def progress_callback(message):
                progress_messages.append({
                    "timestamp": time.time(),
                    "message": message
                })
            
            # Process full document using a single HTML-only generation (no page limit)
            progress_callback("Starting full-document HTML generation...")
            overall_start = time.time()

            html_prompt_text = APP_CONFIG.get('prompts', {}).get('html_generation_prompt', '')
            if not html_prompt_text:
                return Response({"status": "error", "message": "HTML prompt not configured for full document processing"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Call LLM once for full HTML generation
            html_response, error = make_llm_call(
                html_prompt_text,
                absolute_path,
                "text/plain",
                None,  # no page limit for full document
                progress_callback,
                "HTML"
            )

            overall_time = time.time() - overall_start
            progress_callback(f"Full document HTML generation completed in {overall_time:.2f}s")

            if error or not html_response:
                logger.error(f"Failed to generate full-document HTML: {error}")
                return Response({"status": "error", "message": f"Failed to generate HTML: {error}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            if 'candidates' not in html_response:
                logger.error("Invalid HTML response format for full document.", exc_info=True)
                return Response({"status": "error", "message": "Invalid HTML response format"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Extract HTML content
            try:
                html_content = html_response['candidates'][0]['content']['parts'][0]['text']
                html_data = html_content.strip() if html_content and html_content.strip() else None
                pages_processed = html_response.get('pagesProcessed', 1) if isinstance(html_response, dict) else 1
                logger.debug("Successfully generated full-document HTML response")
            except Exception as e:
                logger.error(f"Error extracting HTML from full-document response: {str(e)}", exc_info=True)
                return Response({"status": "error", "message": "Error processing HTML response"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Calculate tokens from HTML usage metadata
            total_input_tokens = 0
            total_output_tokens = 0
            if html_data and html_response and isinstance(html_response, dict) and 'usageMetadata' in html_response:
                html_usage = html_response['usageMetadata']
                total_input_tokens += html_usage.get('promptTokenCount', 0)
                total_output_tokens += html_usage.get('candidatesTokenCount', 0)

            input_tokens = total_input_tokens
            output_tokens = total_output_tokens

            # Update document with full processing results
            # Only HTML view: do not assign json_data
            doc.html_data = html_data
            doc.pages_processed = pages_processed
            doc.is_full_document = True
            doc.input_token = input_tokens
            doc.output_token = output_tokens
            # Ensure api_response_time is defined in this scope
            # Prefer overall_time if available, otherwise try to compute from overall_start
            if 'overall_time' in locals():
                doc.api_response_time = overall_time
            else:
                try:
                    doc.api_response_time = time.time() - overall_start
                except Exception:
                    doc.api_response_time = None
            doc.save()
            
            # Update user usage (difference in pages)
            pages_difference = pages_processed - (doc.pages_processed or 1)
            if pages_difference > 0:
                user.total_pages_processed += pages_difference
                user.save(update_fields=['total_pages_processed'])
            
            logger.info(f"Full document processing completed for document {doc.id}")
            
            return Response({
                "status": "success",
                "message": "Full document processed successfully",
                "pages_processed": pages_processed,
                "progress_messages": progress_messages,
                "usage_info": user.get_usage_info()
            }, status=status.HTTP_200_OK)
            
        except Exception as e:
            logger.error(f"Error in ProcessFullDocumentView: {str(e)}", exc_info=True)
            return Response(
                {"status": "error", "message": f"An error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# New view for generating HTML view of existing documents
class GenerateHTMLView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        document_id = request.data.get("document_id")

        if not document_id:
            return Response(
                {"status": "error", "message": "Missing document_id"},
                status=status.HTTP_400_BAD_REQUEST
            )

        user = request.user

        # Check if user is power user or admin
        if user.user_type not in ['power', 'admin']:
            return Response(
                {"status": "error", "message": "Only power users and admins can generate HTML views"},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            decrypted_id = decrypt_id(document_id)
            doc = get_object_or_404(Document, id=decrypted_id)

            # Check ownership
            is_admin = user.user_type == 'admin'
            if not is_admin and doc.userid_id != user.id:
                return Response(
                    {"error": "You do not have permission to process this document."},
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Check if already has HTML data
            if doc.html_data:
                return Response({
                    "status": "success",
                    "message": "HTML view already exists",
                    "html_data": doc.html_data,
                    "usage_info": user.get_usage_info()
                }, status=status.HTTP_200_OK)

            # Get the original file path
            absolute_path = os.path.join(settings.MEDIA_ROOT, doc.file_path)

            if not os.path.exists(absolute_path):
                return Response(
                    {"status": "error", "message": "Original file not found"},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Determine prompt text for HTML generation
            prompt_text = APP_CONFIG.get('prompts', {}).get('html_generation_prompt', '')

            if not prompt_text:
                return Response(
                    {"status": "error", "message": "Could not determine HTML generation prompt"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            # Progress tracking
            progress_messages = []

            def progress_callback(message):
                progress_messages.append({
                    "timestamp": time.time(),
                    "message": message
                })

            # Determine page limit based on document's current processing
            max_pages = None if doc.is_full_document else 3

            # Generate HTML view using helper function
            response, error = make_llm_call(
                prompt_text,
                absolute_path,
                "text/plain",
                max_pages,
                progress_callback,
                "HTML"
            )

            if error:
                return Response(
                    {"status": "error", "message": f"Failed to generate HTML: {error}"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            if not response or 'candidates' not in response:
                return Response(
                    {"status": "error", "message": "Invalid HTML response from API"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            # Process response
            html_content = response['candidates'][0]['content']['parts'][0]['text']

            if not html_content:
                return Response(
                    {"status": "error", "message": "Empty HTML response from API"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            # Update document with HTML data
            doc.html_data = html_content
            doc.save()

            progress_callback("HTML generation completed successfully")

            return Response({
                "status": "success",
                "message": "HTML view generated successfully",
                "html_data": html_content,
                "progress_messages": progress_messages,
                "usage_info": user.get_usage_info()
            }, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Error in GenerateHTMLView: {str(e)}", exc_info=True)
            return Response(
                {"status": "error", "message": f"An error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# New view for generating PDF from HTML content
class GeneratePDFView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        document_id = request.data.get("document_id")

        if not document_id:
            return Response(
                {"status": "error", "message": "Missing document_id"},
                status=status.HTTP_400_BAD_REQUEST
            )

        user = request.user

        if user.user_type not in ['power', 'admin']:
            return Response(
                {"status": "error", "message": "Only power users and admins can generate PDF documents"},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            decrypted_id = decrypt_id(document_id)
            doc = get_object_or_404(Document, id=decrypted_id)

            if user.user_type != 'admin' and doc.userid_id != user.id:
                return Response(
                    {"error": "You do not have permission to process this document."},
                    status=status.HTTP_403_FORBIDDEN,
                )

            if not doc.html_data:
                return Response(
                    {"status": "error", "message": "No HTML data available for this document."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Generate PDF (no CSS)
            try:
                from weasyprint import HTML
                from weasyprint.text.fonts import FontConfiguration
                from bs4 import BeautifulSoup

                soup = BeautifulSoup(doc.html_data, 'html.parser')
                body = soup.find('body')
                html_content = str(body) if body else doc.html_data

                font_config = FontConfiguration()
                html = HTML(string=html_content)
                pdf_bytes = html.write_pdf(font_config=font_config)

                response = HttpResponse(pdf_bytes, content_type='application/pdf')
                filename = f"document_{decrypted_id}.pdf"
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response

            except (ImportError, OSError):
                # Fallback: ReportLab
               
                soup = BeautifulSoup(doc.html_data, 'html.parser')
                buffer = BytesIO()
                doc_pdf = SimpleDocTemplate(buffer)
                story = []
                styles = getSampleStyleSheet()
                normal = styles['Normal']

                for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'table']):
                    if element.name == 'p':
                        story.append(Paragraph(element.get_text(), normal))
                        story.append(Spacer(1, 12))
                    elif element.name in ['ul', 'ol']:
                        for li in element.find_all('li'):
                            story.append(Paragraph("• " + li.get_text(), normal))
                            story.append(Spacer(1, 6))
                    elif element.name == 'table':
                        rows = element.find_all('tr')
                        data = [[cell.get_text().strip() for cell in row.find_all(['td', 'th'])] for row in rows]
                        table = Table(data)
                        table.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(table)
                        story.append(Spacer(1, 12))

                doc_pdf.build(story)
                pdf_bytes = buffer.getvalue()
                buffer.close()

                response = HttpResponse(pdf_bytes, content_type='application/pdf')
                filename = f"document_{decrypted_id}.pdf"
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response

        except Exception as e:
            logger.error(f"Error in GeneratePDFView: {str(e)}", exc_info=True)
            return Response(
                {"status": "error", "message": f"An error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class GenerateDOCView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        document_id = request.data.get("document_id")

        if not document_id:
            return Response(
                {"status": "error", "message": "Missing document_id"},
                status=status.HTTP_400_BAD_REQUEST
            )

        user = request.user

        if user.user_type not in ['power', 'admin']:
            return Response(
                {"status": "error", "message": "Only power users and admins can generate DOC documents"},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            decrypted_id = decrypt_id(document_id)
            doc = get_object_or_404(Document, id=decrypted_id)

            if user.user_type != 'admin' and doc.userid_id != user.id:
                return Response(
                    {"error": "You do not have permission to process this document."},
                    status=status.HTTP_403_FORBIDDEN,
                )

            if not doc.html_data:
                return Response(
                    {"status": "error", "message": "No HTML data available for this document."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            from docx import Document as DocxDocument
            from bs4 import BeautifulSoup
            from io import BytesIO

            soup = BeautifulSoup(doc.html_data, 'html.parser')
            document = DocxDocument()

            document.add_heading('Document Export', 0)

            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'table']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    document.add_heading(element.get_text(), level)
                elif element.name == 'p':
                    document.add_paragraph(element.get_text())
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        document.add_paragraph(li.get_text(), style='ListBullet')
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        cols = len(rows[0].find_all(['td', 'th']))
                        table = document.add_table(rows=len(rows), cols=cols)
                        table.style = 'Table Grid'
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                table.cell(i, j).text = cell.get_text().strip()

            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"document_{decrypted_id}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        except Exception as e:
            logger.error(f"Error in GenerateDOCView: {str(e)}", exc_info=True)
            return Response(
                {"status": "error", "message": f"An error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )